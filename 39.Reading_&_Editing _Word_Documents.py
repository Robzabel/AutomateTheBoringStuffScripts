import docx, os
os.chdir('/home/zabex/Documents/automate_online-materials/')

d = docx.Document('demo.docx') #call the document to open and save it in to a variable
print(d.paragraphs)# use the parafraphs member variable to call a list of the documents paragraphs
print(d.paragraphs[0].text) # print the text within the first paragraph
p = d.paragraphs[1] #save the text within the second paragraph to a variabel
print(p.runs) #print the runs memory locations
print(p.runs[3].text)#print the text contained in the runs
p.runs[3].underline = True #edit the document so that the 4th run is now underlined
p.runs[3].text = 'Italic and underlined'#change the text of the 4th run
p.runs[3].italic =True #make the 4th run italic
d.save('/home/zabex/Desktop/demo.docx')#save the changes to the document

newDocument = docx.Document() #create a new blank document
newDocument.add_paragraph('Hello this is a paragraph')#add some text to the paragraphs
newDocument.add_paragraph('This is another paragraph')
newDocument.save('/home/zabex/Desktop/demo2.docx')


